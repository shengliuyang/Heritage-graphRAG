# vector.py

import os
import argparse
import ffmpeg
import whisper
from tqdm import tqdm
import re
import jieba
import jieba.analyse
from docx import Document as DocxDocument  # 重命名以避免与 langchain Document 冲突
import pypdf  # 新增：用于处理 PDF 文档
from typing import List, Dict, Any
from collections import Counter
from langchain.schema import Document  # 修改：添加正确的 Document 导入

from langchain_community.document_loaders import TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import OllamaEmbeddings
from langchain.text_splitter import CharacterTextSplitter

# 如果你想使用不同的模型，可以修改这里
WHISPER_MODEL_NAME = "large-v2"
OLLAMA_MODEL_NAME = "mxbai-embed-large:latest"  # 仅用于向量嵌入

def extract_keywords(text: str, topK: int = 5) -> List[str]:
    """
    使用多种方法提取关键词，并进行权重合并
    """
    try:
        # 1. TF-IDF 提取
        tfidf_keywords = set(jieba.analyse.extract_tags(text, topK=topK*2))
        
        # 2. TextRank 提取
        textrank_keywords = set(jieba.analyse.textrank(text, topK=topK*2))
        
        # 3. 合并结果并按重要性排序
        # 同时出现在两种方法中的关键词优先
        common_keywords = list(tfidf_keywords & textrank_keywords)
        other_keywords = list(tfidf_keywords | textrank_keywords - set(common_keywords))
        
        # 4. 按长度过滤（去掉过短的词）
        filtered_keywords = [k for k in (common_keywords + other_keywords) if len(k) > 1]
        
        return filtered_keywords[:topK]
        
    except Exception as e:
        print(f"关键词提取出错: {e}")
        return []

def get_text_summary(text: str, max_length: int = 150) -> str:
    """
    生成更智能的文本摘要
    """
    try:
        # 1. 清理文本
        text = re.sub(r'\s+', ' ', text).strip()
        
        # 2. 分句
        sentences = []
        for sent in re.split(r'([。！？])', text):
            if sent.strip() and sent not in '。！？':
                sentences.append(sent.strip())
        
        if not sentences:
            return text[:max_length]
        
        # 3. 计算每个句子的重要性分数
        sentence_scores = {}
        keywords = extract_keywords(text, topK=10)  # 提取更多关键词用于评分
        
        for sentence in sentences:
            score = 0
            # 根据关键词出现次数评分
            for keyword in keywords:
                score += sentence.count(keyword) * 2
            
            # 根据句子位置评分（开头和结尾的句子更重要）
            if sentence == sentences[0]:
                score += 3
            elif sentence == sentences[-1]:
                score += 2
            
            # 根据句子长度评分（过长或过短的句子降权）
            length = len(sentence)
            if 10 <= length <= 50:
                score += 1
            
            sentence_scores[sentence] = score
        
        # 4. 选择最重要的句子
        sorted_sentences = sorted(sentence_scores.items(), key=lambda x: x[1], reverse=True)
        
        # 5. 组合摘要（保持原有顺序）
        summary_sentences = []
        total_length = 0
        
        for sentence in sentences:
            if sentence in dict(sorted_sentences[:3]):  # 只取前三个最重要的句子
                if total_length + len(sentence) <= max_length:
                    summary_sentences.append(sentence)
                    total_length += len(sentence)
                else:
                    break
        
        summary = '。'.join(summary_sentences)
        
        # 6. 如果摘要太短，直接返回开头部分
        if len(summary) < max_length * 0.5:
            return text[:max_length]
            
        return summary + '。'
        
    except Exception as e:
        print(f"生成摘要时出错: {e}")
        return text[:max_length]

def clean_text(text: str) -> List[str]:
    """
    清理和预处理文本，返回语义完整的句子列表
    """
    # 基础文本清理
    text = re.sub(r'\s+', ' ', text).strip()
    
    # 使用标点符号分割句子
    sentences = []
    # 分割句子，保留分隔符
    parts = re.split(r'([。！？])', text)
    
    # 组合句子和标点
    for i in range(0, len(parts)-1, 2):
        if i+1 < len(parts):
            sentence = parts[i] + parts[i+1]
            if sentence.strip():
                sentences.append(sentence.strip())
    
    # 处理最后一个部分（如果没有标点）
    if len(parts) % 2 == 1 and parts[-1].strip():
        sentences.append(parts[-1].strip())
    
    return sentences

def extract_audio_to_wav(input_video_path, output_audio_path):
    """
    使用 ffmpeg 将视频文件提取为 wav 音频文件
    """
    (
        ffmpeg
        .input(input_video_path)
        .output(output_audio_path, format='wav')
        .run(overwrite_output=True)
    )

def transcribe_audio(whisper_model, audio_path) -> str:
    """
    使用指定的 Whisper 模型对音频进行转录，返回中文文本
    """
    result = whisper_model.transcribe(audio_path, language='zh', fp16=False)
    return result["text"].strip()

def merge_short_chunks(chunks: List[Document], min_length: int = 600) -> List[Document]:
    """
    合并过短的文本片段，确保每个片段至少达到最小长度
    """
    if not chunks:
        return chunks
        
    merged_chunks = []
    current_chunk = None
    
    for chunk in chunks:
        if current_chunk is None:
            current_chunk = chunk
            continue
            
        # 如果当前chunk或下一个chunk小于最小长度，进行合并
        if len(current_chunk.page_content) < min_length or len(chunk.page_content) < min_length:
            # 合并文本内容
            merged_content = current_chunk.page_content + "\n" + chunk.page_content
            
            # 合并并更新元数据
            merged_metadata = current_chunk.metadata.copy()
            merged_metadata.update({
                'content_length': len(merged_content),
                'is_merged': True,
                'original_chunks': merged_metadata.get('original_chunks', 1) + 1
            })
            
            # 重新生成关键词和摘要
            keywords = extract_keywords(merged_content)
            summary = get_text_summary(merged_content)
            merged_metadata['keywords'] = ','.join(keywords)
            merged_metadata['summary'] = summary
            
            # 更新当前chunk
            current_chunk = Document(
                page_content=merged_content,
                metadata=merged_metadata
            )
        else:
            merged_chunks.append(current_chunk)
            current_chunk = chunk
    
    # 添加最后一个chunk
    if current_chunk is not None:
        merged_chunks.append(current_chunk)
    
    return merged_chunks

def process_word_documents(word_folder: str) -> List[Document]:
    """直接处理 Word 文档并返回文档对象列表"""
    if not os.path.exists(word_folder):
        print(f"Word文档文件夹 {word_folder} 不存在")
        return []

    word_files = [f for f in os.listdir(word_folder) if f.lower().endswith(('.doc', '.docx'))]
    processed_docs = []
    
    for word_file in tqdm(word_files, desc="处理Word文档"):
        word_path = os.path.join(word_folder, word_file)
            
        try:
            doc = DocxDocument(word_path)
            
            # 1. 首先收集所有文本内容
            text_content = []
            
            # 处理正文段落
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # 处理表格内容
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        table_text.append(' | '.join(row_text))
                if table_text:
                    text_content.append('\n'.join(table_text))
            
            # 2. 将收集的文本合并成一个完整的文本
            full_text = '\n\n'.join(text_content)
            
            # 3. 使用文本分割器进行分块
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=800,
                chunk_overlap=150,
                separators=["\n\n", "\n", "。", "！", "？", "，", " ", ""],
                length_function=len,
            )
            
            chunks = text_splitter.split_text(full_text)
            
            # 4. 处理每个文本块
            current_chunk = None
            
            for chunk in chunks:
                if not chunk.strip():
                    continue
                    
                metadata = {
                    'source': word_file,
                    'document_type': 'word',
                    'file_name': os.path.splitext(word_file)[0],
                    'content_length': len(chunk),
                    'keywords': ','.join(extract_keywords(chunk)),
                    'summary': get_text_summary(chunk)
                }
                
                doc = Document(page_content=chunk, metadata=metadata)
                
                # 应用合并逻辑
                if current_chunk is None:
                    current_chunk = doc
                elif len(current_chunk.page_content) < 600:
                    # 合并文本
                    merged_content = current_chunk.page_content + "\n" + chunk
                    
                    # 更新元数据
                    merged_metadata = current_chunk.metadata.copy()
                    merged_metadata.update({
                        'content_length': len(merged_content),
                        'is_merged': True,
                        'original_chunks': merged_metadata.get('original_chunks', 1) + 1,
                        'keywords': ','.join(extract_keywords(merged_content)),
                        'summary': get_text_summary(merged_content)
                    })
                    
                    current_chunk = Document(
                        page_content=merged_content,
                        metadata=merged_metadata
                    )
                else:
                    processed_docs.append(current_chunk)
                    current_chunk = doc
            
            # 添加最后一个chunk
            if current_chunk is not None:
                processed_docs.append(current_chunk)
                
        except Exception as e:
            print(f"处理文件 {word_file} 时出错: {str(e)}")
            import traceback
            traceback.print_exc()
            continue

    print(f"Word文档处理完成，共生成 {len(processed_docs)} 个文档片段")
    print(f"其中合并的短文本段落数: {sum(1 for doc in processed_docs if doc.metadata.get('is_merged', False))}")
    return processed_docs

def process_pdf_content(pdf_folder: str) -> List[Dict]:
    """
    直接处理PDF文件内容，返回处理后的文档列表
    """
    if not os.path.exists(pdf_folder):
        print(f"PDF文档文件夹 {pdf_folder} 不存在")
        return []
        
    pdf_files = [f for f in os.listdir(pdf_folder) if f.lower().endswith('.pdf')]
    if not pdf_files:
        return []

    processed_docs = []
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=900,
        chunk_overlap=150,
        separators=["\n\n", "\n", "。", "！", "？", "，", " ", ""],
        length_function=len,
    )
    
    for pdf_file in tqdm(pdf_files, desc="处理PDF文档"):
        pdf_path = os.path.join(pdf_folder, pdf_file)
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                
                # 提取PDF文档的元数据
                pdf_info = pdf_reader.metadata
                title = pdf_info.get('/Title', os.path.splitext(pdf_file)[0])
                author = pdf_info.get('/Author', 'Unknown')
                creation_date = pdf_info.get('/CreationDate', '')
                
                # 处理每一页
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    text = page.extract_text()
                    if not text.strip():
                        continue
                        
                    # 清理和分句
                    cleaned_sentences = clean_text(text)
                    
                    # 对句子进行分块
                    current_chunk = []
                    current_length = 0
                    chunks_text = []
                    
                    for sentence in cleaned_sentences:
                        sentence_length = len(sentence)
                        if current_length + sentence_length > 190000:
                            if current_chunk:
                                chunks_text.append("".join(current_chunk))
                            current_chunk = [sentence]
                            current_length = sentence_length
                        else:
                            current_chunk.append(sentence)
                            current_length += sentence_length
                    
                    if current_chunk:
                        chunks_text.append("".join(current_chunk))
                    
                    # 处理每个文本块
                    for chunk_text in chunks_text:
                        # 提取关键词和生成摘要
                        keywords = extract_keywords(chunk_text)
                        summary = get_text_summary(chunk_text)
                        
                        # 创建增强的元数据 - 将关键词列表转换为字符串
                        metadata = {
                            'source': pdf_file,
                            'document_type': 'pdf',
                            'title': title,
                            'author': author,
                            'creation_date': creation_date,
                            'page_number': page_num,
                            'file_name': os.path.splitext(pdf_file)[0],
                            'content_length': len(chunk_text),
                            'keywords': ','.join(keywords),  # 修改：将关键词列表转换为逗号分隔的字符串
                            'summary': summary,
                        }
                        
                        # 创建文档对象
                        doc = Document(page_content=chunk_text, metadata=metadata)
                        processed_docs.append(doc)
                        
        except Exception as e:
            print(f"处理PDF文件 {pdf_file} 时出错: {str(e)}")
            continue
    
    # 在返回之前合并短文本
    processed_docs = merge_short_chunks(processed_docs, min_length=500)
    
    return processed_docs

def process_txt_files(txt_files: List[str], txt_folder: str) -> List[Document]:
    """处理 txt 文件并返回文档对象列表"""
    all_docs = []
    
    # 配置文本分割器
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=600,  # 主体长度
        chunk_overlap=150,  # 重叠长度
        length_function=len,
        separators=["\n\n", "\n", "。", "！", "？", "，", " ", ""]  # 优先按自然段落分割
    )
    
    for txt_file in tqdm(txt_files, desc="Processing TXT files"):
        try:
            # 读取整个文件内容
            with open(os.path.join(txt_folder, txt_file), 'r', encoding='utf-8') as f:
                text = f.read().strip()
            
            if not text:
                print(f"警告：{txt_file} 内容为空")
                continue
                
            # 使用文本分割器切分文本
            chunks = text_splitter.split_text(text)
            
            # 为每个文本块创建文档对象
            for chunk in chunks:
                # 提取关键词和生成摘要
                keywords = extract_keywords(chunk)
                summary = get_text_summary(chunk)
                
                # 创建元数据
                metadata = {
                    'source': txt_file,
                    'chunk_type': 'transcript',
                    'file_name': os.path.splitext(txt_file)[0],
                    'content_length': len(chunk),
                    'keywords': ','.join(keywords),
                    'summary': summary,
                }
                
                # 创建文档对象
                doc = Document(page_content=chunk, metadata=metadata)
                all_docs.append(doc)
                
        except Exception as e:
            print(f"处理文件 {txt_file} 时出错: {str(e)}")
            continue
            
    print(f"文本处理完成，共生成 {len(all_docs)} 个文档片段")
    return all_docs

def process_mp4_files(
    mp4_folder: str,
    txt_folder: str,
    persist_directory: str,
    word_folder: str = None,
    pdf_folder: str = None
):
    """
    1. 处理所有文档（视频、Word、PDF）
    2. 将所有内容加入到Chroma向量数据库中
    """
    os.makedirs(txt_folder, exist_ok=True)

    # 初始化向量数据库
    embeddings = OllamaEmbeddings(model=OLLAMA_MODEL_NAME)
    vectorstore = Chroma(
        collection_name="txt_documents",
        persist_directory=persist_directory,
        embedding_function=embeddings,
        collection_metadata={
            "hnsw:space": "cosine",
            "hnsw:construction_ef": 400,
            "hnsw:search_ef": 400,
            "hnsw:M": 128,
        }
    )

    # 处理Word文档（如果指定了word_folder）
    if word_folder:
        word_docs = process_word_documents(word_folder)
        if word_docs:
            print(f"\n添加 {len(word_docs)} 个Word文档片段到向量数据库...")
            print(f"其中合并的短文本段落数: {sum(1 for doc in word_docs if doc.metadata.get('is_merged', False))}")
            batch_size = 20
            for i in tqdm(range(0, len(word_docs), batch_size), desc="Adding Word documents"):
                batch_docs = word_docs[i:i + batch_size]
                vectorstore.add_documents(batch_docs)

    # 处理PDF文档
    if pdf_folder:
        pdf_docs = process_pdf_content(pdf_folder)
        if pdf_docs:
            print(f"\n添加 {len(pdf_docs)} 个PDF文档片段到向量数据库...")
            print(f"其中合并的短文本段落数: {sum(1 for doc in pdf_docs if doc.metadata.get('is_merged', False))}")
            batch_size = 20
            for i in tqdm(range(0, len(pdf_docs), batch_size), desc="Adding PDF documents"):
                batch_docs = pdf_docs[i:i + batch_size]
                vectorstore.add_documents(batch_docs)

    # 3) 处理 MP4 文件
    mp4_files = [f for f in os.listdir(mp4_folder) if f.lower().endswith('.mp4')]
    if not mp4_files:
        print("指定文件夹中没有找到 .mp4 文件。")
        return

    # 处理 MP4 文件并生成文本
    for mp4_file in tqdm(mp4_files, desc="Processing MP4 files"):
        mp4_path = os.path.join(mp4_folder, mp4_file)
        base_name = os.path.splitext(mp4_file)[0]
        audio_path = os.path.join(mp4_folder, f"{base_name}.wav")
        txt_path = os.path.join(txt_folder, f"{base_name}.txt")

        if os.path.exists(txt_path):
            print(f"{txt_path} 已存在，跳过转录。")
        else:
            # 提取音频
            extract_audio_to_wav(mp4_path, audio_path)
            # 转文字
            whisper_model = whisper.load_model(WHISPER_MODEL_NAME)
            transcription = transcribe_audio(whisper_model, audio_path)
            # 保存原始文本
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write(transcription)
            # 删除临时音频文件
            os.remove(audio_path)

    # 等待用户输入继续
    input("\n音频转文本完成。按 'c' 继续处理...\n")

    # 4) 处理所有txt文件
    txt_files = [f for f in os.listdir(txt_folder) if f.lower().endswith('.txt')]
    
    # 5) 处理txt文件并获取文档对象
    all_docs = process_txt_files(txt_files, txt_folder)

    # 6) 批量添加文档到向量数据库
    if all_docs:
        batch_size = 20
        for i in tqdm(range(0, len(all_docs), batch_size), desc="Adding to vectorstore"):
            batch_docs = all_docs[i:i + batch_size]
            vectorstore.add_documents(batch_docs)

        print(f"所有文本已成功加入向量数据库，并持久化到 {persist_directory}")
        print(f"总共处理了 {len(all_docs)} 个文档片段")
    else:
        print("没有找到有效的文本段落")

def main():
    parser = argparse.ArgumentParser(description="将视频、Word和PDF文档转换为文本并写入Chroma向量数据库")
    parser.add_argument('--mp4_folder', required=True, help="包含mp4文件的文件夹路径")
    parser.add_argument('--txt_folder', required=True, help="输出txt文件的文件夹路径")
    parser.add_argument('--persist_directory', required=True, help="Chroma向量数据库的持久化目录")
    parser.add_argument('--word_folder', help="包含Word文档的文件夹路径")  # 新增参数
    parser.add_argument('--pdf_folder', help="包含PDF文档的文件夹路径")    # 新增参数

    args = parser.parse_args()

    process_mp4_files(
        mp4_folder=args.mp4_folder,
        txt_folder=args.txt_folder,
        persist_directory=args.persist_directory,
        word_folder=args.word_folder,
        pdf_folder=args.pdf_folder
    )


if __name__ == "__main__":
    main()





